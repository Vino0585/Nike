import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches


class ExecutionReportWriter:
    def __init__(self):
        self.output_dir = Path(__file__).resolve().parents[2] / "Output_files"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.run_id = os.getenv("AU_EXECUTION_RUN_ID", "").strip() or datetime.now().strftime(
            "%Y%m%d_%H%M%S"
        )
        self.output_path = self.output_dir / f"Inbound_Execution_Summary_{self.run_id}.docx"
        self.step_state_path = self.output_dir / f".Inbound_Execution_Steps_{self.run_id}.json"
        self.logo_path = self._resolve_logo_path()

    def _resolve_logo_path(self) -> Path | None:
        env_logo_path = os.getenv("NIKE_LOGO_PATH", "").strip()
        if env_logo_path:
            candidate = Path(env_logo_path)
            if candidate.exists():
                return candidate
        for filename in ("Nike_Logo.png", "Nike_Logo.jpg", "Nike_Logo.jpeg"):
            candidate = self.output_dir / filename
            if candidate.exists():
                return candidate
        return None

    @staticmethod
    def _safe(value):
        if value is None:
            return ""
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=True)
        return str(value)

    def _ensure_doc_header(self, doc: Document):
        if self.logo_path:
            try:
                doc.add_picture(str(self.logo_path), width=Inches(1.8))
            except Exception:
                doc.add_paragraph("NIKE")
        else:
            doc.add_paragraph("NIKE")
        doc.add_heading("Australia Impact Inbound Execution Summary", level=1)
        doc.add_paragraph(f"Run ID: {self.run_id}")
        doc.add_paragraph(f"Document Created At: {datetime.now().isoformat()}")

    def _load_step_state(self) -> list[dict]:
        if not self.step_state_path.exists():
            return []
        try:
            raw = json.loads(self.step_state_path.read_text(encoding="utf-8"))
            if isinstance(raw, list):
                return raw
            return []
        except Exception:
            return []

    def _save_step_state(self, state: list[dict]):
        self.step_state_path.write_text(json.dumps(state, ensure_ascii=True, indent=2), encoding="utf-8")

    def write_step_report(
        self,
        step_name: str,
        run_user: str,
        started_at: datetime,
        ended_at: datetime,
        status: str,
        summary: dict,
        records: list[dict],
    ) -> Path:
        duration_seconds = (ended_at - started_at).total_seconds()
        state = self._load_step_state()
        state.append(
            {
                "step_name": self._safe(step_name),
                "run_user": self._safe(run_user or "unknown"),
                "started_at": started_at.isoformat(),
                "ended_at": ended_at.isoformat(),
                "duration_seconds": f"{duration_seconds:.2f}",
                "status": self._safe(status),
                "summary": summary or {},
                "records": records or [],
            }
        )
        self._save_step_state(state)
        return self.output_path

    def write_end_to_end_summary(
        self,
        run_user: str,
        started_at: datetime,
        ended_at: datetime,
        status: str,
        selected_flags: dict,
        completed_steps: list[str],
        error_message: str = "",
    ) -> Path:
        total_seconds = (ended_at - started_at).total_seconds()
        step_state = self._load_step_state()
        doc = Document()
        self._ensure_doc_header(doc)

        doc.add_heading("End-to-End Execution Summary", level=2)
        meta_table = doc.add_table(rows=0, cols=2)
        for key, value in (
            ("Overall Status", status),
            ("Run User", run_user or "unknown"),
            ("Execution Started At", started_at.isoformat()),
            ("Execution Ended At", ended_at.isoformat()),
            ("Total Execution Time (seconds)", f"{total_seconds:.2f}"),
            ("Completed Steps", ", ".join(completed_steps) if completed_steps else "None"),
        ):
            row = meta_table.add_row().cells
            row[0].text = self._safe(key)
            row[1].text = self._safe(value)

        doc.add_paragraph("Selected Flags")
        if selected_flags:
            for key, value in selected_flags.items():
                doc.add_paragraph(f"{self._safe(key)}: {self._safe(value)}", style="List Bullet")
        else:
            doc.add_paragraph("No flags found.", style="List Bullet")

        if error_message:
            doc.add_paragraph("Failure Reason")
            doc.add_paragraph(self._safe(error_message), style="List Bullet")

        doc.add_paragraph(
            "Detailed execution summary for each steps are available from Page 2 onwards."
        )
        doc.add_page_break()

        for step in step_state:
            doc.add_heading(self._safe(step.get("step_name", "Step")), level=2)
            step_meta = doc.add_table(rows=0, cols=2)
            for key, value in (
                ("Status", step.get("status", "")),
                ("Run User", step.get("run_user", "unknown")),
                ("Started At", step.get("started_at", "")),
                ("Ended At", step.get("ended_at", "")),
                ("Execution Time (seconds)", step.get("duration_seconds", "")),
            ):
                row = step_meta.add_row().cells
                row[0].text = self._safe(key)
                row[1].text = self._safe(value)

            doc.add_paragraph("Summary")
            step_summary = step.get("summary", {}) if isinstance(step.get("summary", {}), dict) else {}
            if step_summary:
                for key, value in step_summary.items():
                    doc.add_paragraph(f"{self._safe(key)}: {self._safe(value)}", style="List Bullet")
            else:
                doc.add_paragraph("No summary details captured.", style="List Bullet")

            doc.add_paragraph("Execution Details")
            step_records = step.get("records", []) if isinstance(step.get("records", []), list) else []
            if step_records:
                for index, record in enumerate(step_records, start=1):
                    doc.add_paragraph(f"Record {index}")
                    details_table = doc.add_table(rows=0, cols=2)
                    if isinstance(record, dict):
                        for key, value in record.items():
                            row = details_table.add_row().cells
                            row[0].text = self._safe(key)
                            row[1].text = self._safe(value)
            else:
                doc.add_paragraph("No step records captured.", style="List Bullet")

            doc.add_page_break()

        doc.save(self.output_path)
        try:
            if self.step_state_path.exists():
                self.step_state_path.unlink()
        except Exception:
            pass
        return self.output_path
